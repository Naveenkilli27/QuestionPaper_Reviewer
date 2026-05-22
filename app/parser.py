"""Extract plain text from uploaded question papers (PDF, DOCX, TXT)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def _utf8_text(s: str) -> str:
    return s.encode("utf-8", errors="replace").decode("utf-8")


def _normalize_pdf_spaces(text: str) -> str:
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text


def _marks_suffix_from_cells(cells: list[str]) -> str:
    """If a table cell looks like a marks-only value, tag it for the AI (typical QP col 3)."""
    if len(cells) <= 3:
        return ""
    t = cells[3].strip()
    if not t or len(t) > 24:
        return ""
    m = re.match(r"^\(?\s*(\d{1,2})\s*\)?\s*$", t)
    if m:
        return f"[M:{m.group(1)}]"
    m2 = re.match(r"^(\d{1,2})\s*marks?$", t, re.IGNORECASE)
    if m2:
        return f"[M:{m2.group(1)}]"
    return ""


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _utf8_text(_normalize_pdf_spaces(_pdf_text(path)))
    if suffix == ".docx":
        return _utf8_text(_docx_text(path))
    if suffix in {".txt", ".md"}:
        return _utf8_text(path.read_text(encoding="utf-8", errors="replace"))
    raise ValueError(f"Unsupported file type: {suffix}. Use PDF, DOCX, or TXT.")


def _pdf_text(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ValueError(f"Cannot open PDF (corrupt or unsupported): {exc}") from exc
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception as exc:
            t = f"\n[Page extract error: {exc}]\n"
        parts.append(t)
    return "\n".join(parts)


def _docx_text(path: Path) -> str:
    """
    Extract all questions from a DOCX question paper.

    University papers use a table layout:
      Col 0: Q number  ("6.a", "b", "7.a" or blank for Part A)
      Col 1: Part      ("(i)", "(ii)" or blank)
      Col 2: Question text
      Col 3: Marks box "(  )"
      Col 4: CO        "CO1"
      Col 5: KL/Bloom  "KL2"

    Part A questions have NO question number in col 0 — we auto-number them Q1..Q5.
    """
    doc = Document(str(path))
    lines: list[str] = []

    # Header paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue

        ncols = len(rows[0].cells)

        # Detect university QP table: 5-6 cols, CO values in col 4
        is_qp_table = False
        if ncols >= 5:
            sample: list[str] = []
            for row in list(rows)[1:8]:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) > 4:
                    sample.append(cells[4])
            if any(re.match(r"CO\d+", v, re.IGNORECASE) for v in sample):
                is_qp_table = True

        if not is_qp_table:
            # Generic table fallback
            for row in rows:
                cell_texts = [c.text.strip() for c in row.cells if c.text.strip()]
                deduped: list[str] = []
                for ct in cell_texts:
                    if not deduped or ct != deduped[-1]:
                        deduped.append(ct)
                if deduped:
                    lines.append(" | ".join(deduped))
            continue

        # ---- University QP table parsing ----
        current_section = ""   # PART A / PART B / PART C
        current_qnum   = ""    # e.g. "6.a"
        part_a_counter = 0     # auto-number Part A questions Q1..Q5

        for row in rows:
            cells = [c.text.strip() for c in row.cells]
            while len(cells) < 6:
                cells.append("")

            col0   = cells[0]   # Q number or section header
            col1   = cells[1]   # part (i)(ii) or blank
            # col 2 and col 3 are often duplicates due to merged cells;
            # use the longer of the two as the question text
            qtext  = cells[2] if len(cells[2]) >= len(cells[3]) else cells[3]
            co_val = cells[4]
            kl_val = cells[5]

            # --- Section header row (Part A/B/C, Section/Unit headings) ---
            _sec = col0.strip()
            if re.match(r"^PART\s+", _sec, re.IGNORECASE):
                current_section = _sec.split("(")[0].strip()
                current_qnum = ""
                continue
            if re.match(r"^(SECTION|UNIT)\b", _sec, re.IGNORECASE) and len(_sec) < 120:
                current_section = _sec
                current_qnum = ""
                continue
            # Section title sometimes only in col1/col2 (merged layout)
            if not _sec and col1:
                _cand = col1.strip()
                if re.match(r"^(PART|SECTION|UNIT)\b", _cand, re.IGNORECASE) and len(_cand) < 120:
                    current_section = _cand.split("(")[0].strip()
                    current_qnum = ""
                    continue

            # --- OR separator ---
            if re.match(r"^\(OR\)$", col0, re.IGNORECASE):
                lines.append("OR")
                continue

            # --- Divider / empty ---
            if re.match(r"^\*+$", col0) or (not col0 and not qtext):
                continue

            # --- Part A: questions with no number in col 0 ---
            _in_part_a = bool(re.match(r"PART\s*[-–]?\s*A\b", current_section, re.IGNORECASE))
            if not col0 and _in_part_a:
                if qtext and len(qtext.strip()) >= 3 and qtext != col0:
                    part_a_counter += 1
                    qid = f"Q{part_a_counter}"
                    co_tag = f"[{co_val}]" if re.match(r"CO\d+", co_val, re.IGNORECASE) else ""
                    kl_tag = f"[{kl_val}]" if re.match(r"KL?\d+", kl_val, re.IGNORECASE) else ""
                    m_note = _marks_suffix_from_cells(cells)
                    lines.append(f"{qid} | {qtext} {m_note} {co_tag} {kl_tag}".strip())
                continue

            # --- Part B / C: numbered questions ---
            if col0 and col0 not in ("CO", "KL", "(OR)"):
                # Normalize Q number: "6.a" -> "Q6a", "b" -> append to current main Q
                raw_qnum = re.sub(r"\s+", "", col0)
                if re.match(r"^\d", raw_qnum):
                    # Starts with digit — main question like "6.a"
                    norm = "Q" + raw_qnum.replace(".", "")
                    current_qnum = norm
                elif raw_qnum.lower() in ("b", "c", "d"):
                    # Sub option like "b" for OR alternative — keep context
                    # e.g. "6b" if current was "6a"
                    base = re.sub(r"[aA]$", "", current_qnum)
                    current_qnum = base + raw_qnum.upper()
                else:
                    current_qnum = raw_qnum if raw_qnum.upper().startswith("Q") else "Q" + raw_qnum

            # Build part label
            part_label = col1 if col1 else ""

            if not qtext or len(qtext.strip()) < 2:
                continue

            # Deduplicate qtext (merged cells sometimes repeat it)
            if qtext == cells[3]:
                pass  # fine, just one copy

            co_tag = f"[{co_val}]" if re.match(r"CO\d+", co_val, re.IGNORECASE) else ""
            kl_tag = f"[{kl_val}]" if re.match(r"KL?\d+", kl_val, re.IGNORECASE) else ""
            m_note = _marks_suffix_from_cells(cells)

            prefix = f"{current_qnum} {part_label}".strip()
            lines.append(f"{prefix} | {qtext} {m_note} {co_tag} {kl_tag}".strip())

    return "\n".join(lines)
